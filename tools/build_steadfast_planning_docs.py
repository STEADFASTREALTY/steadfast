from pathlib import Path
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_steadfast_listing_workflow as wf


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
TEMPLATE = DOCS / "SteadFast_MVP_Product_Requirements_v0.1.docx"

CONFIGS = [
    {
        "source": DOCS / "SteadFast_Application_and_API_Architecture_v0.1.md",
        "output": DOCS / "SteadFast_Application_and_API_Architecture_v0.1.docx",
        "title": "SteadFast Application and API Architecture",
        "subtitle": "Portable, brokerage-isolated SaaS architecture for the Jamaica MVP",
        "kicker": "TECHNICAL ARCHITECTURE SPECIFICATION",
        "header": "APPLICATION AND API ARCHITECTURE",
        "status": "Target architecture; implementation has not yet begun",
        "diagram": "architecture",
    },
    {
        "source": DOCS / "SteadFast_UX_Sitemap_and_User_Flows_v0.1.md",
        "output": DOCS / "SteadFast_UX_Sitemap_and_User_Flows_v0.1.docx",
        "title": "SteadFast UX Sitemap and User Flows",
        "subtitle": "No-training interaction baseline for public, professional and platform users",
        "kicker": "PRODUCT EXPERIENCE SPECIFICATION",
        "header": "UX SITEMAP AND USER FLOWS",
        "status": "MVP interaction baseline",
        "diagram": "ux",
    },
    {
        "source": DOCS / "steadfast-threat-model.md",
        "output": DOCS / "SteadFast_Security_Plan_and_Threat_Model_v0.1.docx",
        "title": "SteadFast Security Plan and Threat Model",
        "subtitle": "Threat-led controls for an internet-facing multi-brokerage platform",
        "kicker": "SECURITY AND RISK SPECIFICATION",
        "header": "SECURITY PLAN AND THREAT MODEL",
        "status": "Pre-implementation security baseline",
        "diagram": "security",
    },
    {
        "source": DOCS / "SteadFast_Testing_and_Deployment_Plan_v0.1.md",
        "output": DOCS / "SteadFast_Testing_and_Deployment_Plan_v0.1.docx",
        "title": "SteadFast Testing and Deployment Plan",
        "subtitle": "Continuous previews, isolated data, release gates and recoverable production",
        "kicker": "QUALITY AND RELEASE SPECIFICATION",
        "header": "TESTING AND DEPLOYMENT PLAN",
        "status": "MVP delivery baseline",
    },
    {
        "source": DOCS / "SteadFast_Development_Backlog_v0.1.md",
        "output": DOCS / "SteadFast_Development_Backlog_v0.1.docx",
        "title": "SteadFast Development Backlog",
        "subtitle": "Ordered, security-aware implementation plan from foundation through pilot",
        "kicker": "MVP DELIVERY BACKLOG",
        "header": "DEVELOPMENT BACKLOG",
        "status": "Ordered MVP implementation backlog",
    },
    {
        "source": DOCS / "SteadFast_Pilot_and_Legal_Readiness_Pack_v0.1.md",
        "output": DOCS / "SteadFast_Pilot_and_Legal_Readiness_Pack_v0.1.docx",
        "title": "SteadFast Pilot and Legal Readiness Pack",
        "subtitle": "Business and technical draft for Jamaican counsel review",
        "kicker": "PILOT AND LEGAL DRAFT",
        "header": "PILOT AND LEGAL READINESS",
        "status": "Not legal advice; counsel approval required before use",
    },
]


def font(size, bold=False):
    for path in (
        Path(r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
    ):
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def diagram_nodes(kind):
    if kind == "architecture":
        return (
            "SteadFast System Context",
            [
                ("Visitors", 65, 205, "#DCEBFA"),
                ("Professional\nworkspaces", 65, 405, "#DCEBFA"),
                ("Next.js on Vercel", 590, 300, "#FFF0D8"),
                ("Supabase Auth\nPostgreSQL • Storage", 1115, 185, "#E5F3E8"),
                ("Outbox and workers", 1115, 405, "#ECE7F6"),
                ("Email • Maps • Payments\nApproved listing channels", 1570, 300, "#E9EDF2"),
            ],
            [(0, 2), (1, 2), (2, 3), (2, 4), (4, 5)],
        )
    if kind == "ux":
        return (
            "Capability-Aware Experience Map",
            [
                ("Public search\nand property", 65, 285, "#DCEBFA"),
                ("Sign in", 455, 285, "#FFF0D8"),
                ("Agent\nworkspace", 820, 110, "#E5F3E8"),
                ("Approval\ndesk", 820, 285, "#E5F3E8"),
                ("Brokerage\ncontrol", 820, 460, "#E5F3E8"),
                ("Operations", 1300, 200, "#ECE7F6"),
                ("Administration", 1300, 390, "#ECE7F6"),
                ("Approved public\nlisting", 1690, 285, "#DCEBFA"),
            ],
            [(0, 1), (1, 2), (1, 3), (1, 4), (1, 5), (1, 6), (2, 3), (3, 7), (4, 7)],
        )
    return (
        "SteadFast Trust Boundaries",
        [
            ("Untrusted\nbrowser", 65, 285, "#FDE2E2"),
            ("Vercel and\nNext.js", 485, 285, "#FFF0D8"),
            ("Authorization\nand domain rules", 905, 285, "#DCEBFA"),
            ("Supabase\nsystem of record", 1325, 175, "#E5F3E8"),
            ("Private media\nand outbox", 1325, 405, "#ECE7F6"),
            ("External\nproviders", 1740, 285, "#E9EDF2"),
        ],
        [(0, 1), (1, 2), (2, 3), (2, 4), (4, 5)],
    )


def build_diagram(kind):
    ASSETS.mkdir(exist_ok=True)
    output = ASSETS / f"SteadFast_{kind.title()}_Diagram_v0.1.png"
    width, height = 2100, 700
    image = Image.new("RGB", (width, height), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    title, nodes, edges = diagram_nodes(kind)
    draw.text((65, 42), title, font=font(38, True), fill="#17365D")
    box_w, box_h = 300, 130
    centers = [(x + box_w / 2, y + box_h / 2) for _, x, y, _ in nodes]
    for a, b in edges:
        x1, y1 = centers[a]
        x2, y2 = centers[b]
        draw.line((x1, y1, x2, y2), fill="#8793A1", width=5)
    for label, x, y, fill in nodes:
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill=fill, outline="#1F4E78", width=4)
        lines = label.split("\n")
        for idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font(23, True))
            tw = bbox[2] - bbox[0]
            draw.text((x + (box_w - tw) / 2, y + 35 + idx * 31), line, font=font(23, True), fill="#17365D")
    draw.text((65, 645), "Arrows show principal request, data, approval or delivery direction. Detailed controls are defined in the document.", font=font(18), fill="#667085")
    image.save(output, "PNG", optimize=True)
    return output


def configure(doc, config):
    wf.configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    wf.set_font(header.add_run(f"STEADFAST  |  {config['header']}"), size=8.5, bold=True, color=wf.MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    wf.set_font(footer.add_run("Planning Baseline  |  Version 0.1  |  Page "), size=9, color=wf.MID_GRAY)
    wf.add_page_number(footer)


def add_masthead(doc, config):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    wf.set_font(p.add_run(config["kicker"]), size=10, bold=True, color=wf.BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    wf.set_font(p.add_run(config["title"]), size=27, bold=True, color=wf.DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    wf.set_font(p.add_run(config["subtitle"]), size=13.5, color=wf.MID_GRAY)
    for label, value in (
        ("Document owner", "SteadFast Product and Engineering"),
        ("Status", config["status"]),
        ("Version", "0.1"),
        ("Date", "16 July 2026"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        wf.set_font(p.add_run(f"{label}: "), size=10.5, bold=True, color=wf.DARK_BLUE)
        wf.set_font(p.add_run(value), size=10.5, color=wf.DARK)
    doc.add_paragraph()


def add_picture(doc, path, kind):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = p.add_run().add_picture(str(path), width=Inches(6.5))
    inline._inline.docPr.set("name", f"SteadFast {kind} diagram")
    inline._inline.docPr.set("descr", f"High-level SteadFast {kind} diagram. Detailed relationships are explained in the surrounding document.")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    wf.set_font(caption.add_run("Figure 1. High-level relationship map."), size=9, italic=True, color=wf.MID_GRAY)


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    wf.set_font(p.add_run(text), size=10.5, italic=True, color=wf.DARK_BLUE)


def build_one(config):
    doc = Document(TEMPLATE)
    wf.clear_body(doc)
    configure(doc, config)
    add_masthead(doc, config)
    diagram = build_diagram(config["diagram"]) if config.get("diagram") else None

    lines = config["source"].read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("**Version:") or line.startswith("**Prepared:") or line.startswith("**Status:") or line.startswith("**Scope:") or line.startswith("**Applies to:") or line.startswith("**Platforms:") or line.startswith("**Priority:") or line.startswith("**Design goal:") or line.startswith("**Important:"):
            i += 1
            continue
        if line == "~~~mermaid":
            i += 1
            while i < len(lines) and lines[i].strip() != "~~~":
                i += 1
            i += 1
            if diagram:
                add_picture(doc, diagram, config["diagram"])
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = wf.parse_table(lines, i)
            wf.add_table(doc, rows)
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            wf.add_inline(p, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            wf.add_inline(p, line[2:])
            i += 1
            continue
        if line.startswith("> "):
            add_quote(doc, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        wf.add_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = config["title"] + " v0.1"
    props.subject = config["subtitle"]
    props.author = "SteadFast"
    props.keywords = "SteadFast, SaaS, real estate, Jamaica, planning"
    doc.save(config["output"])
    print(config["output"])


def build():
    for config in CONFIGS:
        build_one(config)


if __name__ == "__main__":
    build()
