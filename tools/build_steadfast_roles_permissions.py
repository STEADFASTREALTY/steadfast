from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SteadFast_Roles_and_Permissions_Matrix_v0.1.md"
OUTPUT = ROOT / "docs" / "SteadFast_Roles_and_Permissions_Matrix_v0.1.docx"

BLUE = "2367A6"
DARK_BLUE = "163A5F"
DARK = "1F2933"
MID_GRAY = "5F6B76"
LIGHT_BLUE = "E8F0F7"
PALE_BLUE = "F5F8FB"
WHITE = "FFFFFF"
TABLE_WIDTH = 13992
TABLE_INDENT = 90


def set_font(run, *, size=9.5, bold=False, italic=False, color=DARK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, end])
    set_font(run, size=8.5, color=MID_GRAY)


def add_inline_markdown(paragraph, text, *, size=9.5, color=DARK):
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`") and token.endswith("`"):
            set_font(paragraph.add_run(token[1:-1]), size=max(8, size - 0.5), color=DARK_BLUE, name="Consolas")
        else:
            set_font(paragraph.add_run(token), size=size, color=color)


def column_widths(headers):
    count = len(headers)
    if count == 8:
        first = 4200
        remainder = TABLE_WIDTH - first
        base = remainder // 7
        widths = [first] + [base] * 6 + [remainder - base * 6]
    elif count == 2:
        widths = [3600, TABLE_WIDTH - 3600]
    elif count == 3:
        widths = [3000, 5000, TABLE_WIDTH - 8000]
    else:
        base = TABLE_WIDTH // count
        widths = [base] * (count - 1) + [TABLE_WIDTH - base * (count - 1)]
    return widths


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    widths = column_widths(headers)
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_repeat_header(table.rows[0])
    for r_idx, data in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        for c_idx, value in enumerate(data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_shading(cell, BLUE)
                add_inline_markdown(p, value, size=8.2 if len(headers) == 8 else 9, color=WHITE)
                for run in p.runs:
                    run.bold = True
            else:
                if len(headers) == 8 and c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    size = 7.6
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    size = 8.5
                if r_idx % 2 == 0:
                    set_cell_shading(cell, PALE_BLUE)
                add_inline_markdown(p, value, size=size)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        parts = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            rows.append(parts)
        index += 1
    return rows, index


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 13, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 10.5, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("STEADFAST  |  ROLES AND PERMISSIONS MATRIX"), size=8.3, bold=True, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Authorization Baseline  |  Version 0.1  |  Page "), size=8.5, color=MID_GRAY)
    add_page_number(footer)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("# ") and first_title:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(2)
            set_font(p.add_run("AUTHORIZATION SPECIFICATION"), size=9.5, bold=True, color=BLUE)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            set_font(p.add_run(line[2:]), size=27, bold=True, color=DARK_BLUE)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(14)
            set_font(p.add_run("Brokerage-rooted access control for the Jamaica MVP"), size=13, color=MID_GRAY)
            first_title = False
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(3)
            add_inline_markdown(p, text)
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(3)
            add_inline_markdown(p, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline_markdown(p, line)
        i += 1

    core_props = doc.core_properties
    core_props.title = "SteadFast Roles and Permissions Matrix v0.1"
    core_props.subject = "MVP authorization and brokerage isolation baseline"
    core_props.author = "SteadFast"
    core_props.keywords = "SteadFast, roles, permissions, authorization, brokerage, listings"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
